import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from groq import Groq
from langdetect import detect, DetectorFactory
import re
import tempfile
import os
from typing import List, Tuple
import time

# Fix randomness in language detection
DetectorFactory.seed = 0

# -------------------------------
# Helper functions
# -------------------------------

def detect_language(text: str) -> str:
    """Detect language of a text sample."""
    try:
        lang = detect(text)
        # Map ISO codes to human-readable names
        lang_map = {
            'ar': 'Arabic', 'en': 'English', 'fr': 'French', 'es': 'Spanish',
            'de': 'German', 'it': 'Italian', 'pt': 'Portuguese', 'ru': 'Russian',
            'zh-cn': 'Chinese (Simplified)', 'ja': 'Japanese', 'ko': 'Korean'
        }
        return lang_map.get(lang, lang.upper())
    except:
        return "Unknown"

def extract_document_stats(doc) -> dict:
    """Extract basic stats from a python-docx Document."""
    full_text = []
    table_count = len(doc.tables)
    picture_count = 0
    
    # Count pictures (inline shapes)
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            if run._element.xpath('.//a:blip'):
                picture_count += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run._element.xpath('.//a:blip'):
                            picture_count += 1
    
    # Extract text for length/words
    for paragraph in doc.paragraphs:
        if paragraph.text:
            full_text.append(paragraph.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if paragraph.text:
                        full_text.append(paragraph.text)
    
    text_str = ' '.join(full_text)
    char_count = len(text_str)
    word_count = len(text_str.split())
    approx_pages = max(1, word_count // 300)  # rough estimate
    
    return {
        "characters": char_count,
        "words": word_count,
        "approx_pages": approx_pages,
        "tables": table_count,
        "pictures": picture_count,
        "preview": text_str[:1000] + ("..." if len(text_str) > 1000 else "")
    }

def translate_text(client: Groq, text: str, source_lang: str, target_lang: str, model: str) -> str:
    """Translate a single text chunk using Groq API."""
    if not text or not text.strip():
        return text
    
    prompt = f"""Translate the following text from {source_lang} to {target_lang}.
Preserve numbers, punctuation, special characters, and formatting markers like spaces and newlines.
Do not add any extra commentary or notes. Only output the translated text.

Text:
{text}

Translated text:"""
    
    try:
        response = client.chat.completions.create(
            messages=[{"role": "user", "content": prompt}],
            model=model,
            temperature=0.3,
        )
        translated = response.choices[0].message.content
        # Remove any possible leading/trailing quotes
        translated = translated.strip()
        return translated
    except Exception as e:
        st.error(f"Translation error: {str(e)}")
        return text  # fallback to original

def process_paragraphs_in_chunks(paragraphs: List, client: Groq, source_lang: str, target_lang: str, model: str, progress_bar, chunk_size=20):
    """
    Process a list of paragraph objects in chunks to avoid overloading the API.
    Returns a list of translated texts in the same order.
    """
    total = len(paragraphs)
    translated_texts = [""] * total
    
    for i in range(0, total, chunk_size):
        chunk = paragraphs[i:i+chunk_size]
        chunk_indices = list(range(i, min(i+chunk_size, total)))
        
        # Prepare batch: translate each paragraph individually (parallelism not needed, but sequential safe)
        for idx, para in zip(chunk_indices, chunk):
            original_text = para.text
            if original_text.strip():
                translated = translate_text(client, original_text, source_lang, target_lang, model)
                translated_texts[idx] = translated
            else:
                translated_texts[idx] = original_text
            
            # Update progress
            progress_bar.progress((idx+1)/total)
    
    return translated_texts

def apply_translated_text_to_paragraph(paragraph, translated_text, original_paragraph):
    """
    Replace paragraph text while preserving basic formatting and alignment.
    """
    # Store original style and alignment
    style = paragraph.style
    alignment = paragraph.paragraph_format.alignment
    
    # Clear existing runs
    paragraph.clear()
    
    # Add new run with translated text
    run = paragraph.add_run(translated_text)
    
    # Copy formatting from the first run of the original paragraph (if any)
    if original_paragraph.runs:
        first_run = original_paragraph.runs[0]
        if first_run.bold:
            run.bold = True
        if first_run.italic:
            run.italic = True
        if first_run.underline:
            run.underline = True
        if first_run.font.size:
            run.font.size = first_run.font.size
        if first_run.font.name:
            run.font.name = first_run.font.name
        if first_run.font.color.rgb:
            run.font.color.rgb = first_run.font.color.rgb
    else:
        # Use paragraph style defaults
        run.font.size = Pt(11)
    
    # Restore alignment
    paragraph.alignment = alignment
    # Restore style
    paragraph.style = style

def translate_document(doc, client, source_lang, target_lang, model, progress_callback=None):
    """Translate entire document preserving structure (tables, headers, footers)."""
    # Collect all text elements: paragraphs in body, tables, headers, footers
    elements = []  # list of (container, element_type, original_paragraph)
    
    # Body paragraphs
    for para in doc.paragraphs:
        elements.append(('body', para, para))
    
    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    elements.append(('table_cell', para, para))
    
    # Headers
    for section in doc.sections:
        for para in section.header.paragraphs:
            elements.append(('header', para, para))
        for para in section.first_page_header.paragraphs:
            elements.append(('header', para, para))
        for para in section.even_page_header.paragraphs:
            elements.append(('header', para, para))
    
    # Footers
    for section in doc.sections:
        for para in section.footer.paragraphs:
            elements.append(('footer', para, para))
        for para in section.first_page_footer.paragraphs:
            elements.append(('footer', para, para))
        for para in section.even_page_footer.paragraphs:
            elements.append(('footer', para, para))
    
    total = len(elements)
    if progress_callback:
        progress_bar = st.progress(0)
    
    # Process in chunks of 15-20 paragraphs to avoid rate limits
    chunk_size = 15
    for i in range(0, total, chunk_size):
        chunk_elements = elements[i:i+chunk_size]
        # Translate texts
        for idx, (container, para, original_para) in enumerate(chunk_elements):
            if para.text.strip():
                translated = translate_text(client, para.text, source_lang, target_lang, model)
                # Apply back to the paragraph
                apply_translated_text_to_paragraph(para, translated, original_para)
            if progress_callback:
                progress_bar.progress((i+idx+1)/total)
    
    if progress_callback:
        progress_bar.empty()
    
    return doc

# -------------------------------
# Streamlit UI
# -------------------------------

st.set_page_config(page_title="Doc Translator with Groq", layout="wide")
st.title("📄 Smart Document Translator")
st.markdown("Upload a Word document, detect its language, and translate it while preserving formatting.")

# Sidebar for configuration
with st.sidebar:
    st.header("⚙️ Configuration")
    groq_api_key = st.text_input("Groq API Key", type="password", help="Enter your Groq API key")
    model_name = st.text_input("Model Name", value="openai/gpt-oss-120b", help="Groq model (e.g., llama3-70b-8192, mixtral-8x7b-32768)")
    st.markdown("---")
    st.info("💡 You can get a Groq API key from [console.groq.com](https://console.groq.com)")

# Main area
uploaded_file = st.file_uploader("Choose a Word document (.docx)", type=["docx"])

if uploaded_file is not None:
    # Save uploaded file temporarily
    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp_file:
        tmp_file.write(uploaded_file.getvalue())
        tmp_path = tmp_file.name
    
    # Load document
    doc = Document(tmp_path)
    
    # Extract stats and preview
    stats = extract_document_stats(doc)
    
    # Detect language from the first ~2000 characters
    sample_text = stats['preview'][:2000]
    detected_lang = detect_language(sample_text)
    
    col1, col2 = st.columns(2)
    with col1:
        st.subheader("📊 Document Statistics")
        st.write(f"**Detected language:** {detected_lang}")
        st.write(f"**Characters:** {stats['characters']:,}")
        st.write(f"**Words:** {stats['words']:,}")
        st.write(f"**Approx. pages:** {stats['approx_pages']}")
        st.write(f"**Tables:** {stats['tables']}")
        st.write(f"**Inline pictures:** {stats['pictures']}")
    
    with col2:
        st.subheader("📄 Text Preview")
        st.text_area("First 1000 characters", stats['preview'], height=200)
    
    # Target language selection
    target_langs = ["Arabic", "English", "French", "Spanish", "German", "Italian", "Portuguese", "Russian", "Chinese (Simplified)", "Japanese", "Korean"]
    target_lang = st.selectbox("Translate to:", target_langs, index=0)
    
    # Source language (can be overridden if detection is wrong)
    source_lang = st.text_input("Source language (detected above, edit if needed):", value=detected_lang)
    
    if st.button("🌍 Translate Document", type="primary"):
        if not groq_api_key:
            st.error("Please enter your Groq API key in the sidebar.")
        else:
            # Initialize Groq client
            client = Groq(api_key=groq_api_key)
            
            # Confirm with user about processing time
            st.info(f"Translation will process {stats['words']} words in chunks. This may take a moment...")
            
            # Show progress
            progress_placeholder = st.empty()
            with st.spinner("Translating... please wait."):
                try:
                    translated_doc = translate_document(
                        doc, client, source_lang, target_lang, model_name,
                        progress_callback=True
                    )
                    st.success("✅ Translation completed!")
                    
                    # Save translated document
                    out_path = tmp_path.replace(".docx", "_translated.docx")
                    translated_doc.save(out_path)
                    
                    with open(out_path, "rb") as f:
                        st.download_button(
                            label="📥 Download Translated Document",
                            data=f,
                            file_name=f"translated_{target_lang}.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                except Exception as e:
                    st.error(f"Translation failed: {str(e)}")
                finally:
                    # Cleanup
                    if os.path.exists(tmp_path):
                        os.unlink(tmp_path)
                    if os.path.exists(out_path):
                        os.unlink(out_path)
    
    # Cleanup temp file if not used
    if not st.button("Translate"):
        os.unlink(tmp_path)